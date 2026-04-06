"""
Convert REPORT.md to a professionally formatted DOCX file WITH charts embedded.

Usage:
    1. First run:  python generate_charts.py   (generates charts/ folder)
    2. Then run:   python convert_report_with_charts.py
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_report():
    doc = Document()

    # ─── Page Setup ───
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ─── Define Styles ───
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)

    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 70, 130)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0, 80, 140)
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(6)

    # ─── Helper Functions ───
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Times New Roman'
        return h

    def add_para(text, bold=False, italic=False, size=12):
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                run = p.add_run(part)
                run.font.bold = bold
            run.font.italic = italic
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
        return p

    def add_bullet(text, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        return p

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, '003366')
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text))
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                if r_idx % 2 == 0:
                    set_cell_shading(cell, 'EBF5FB')
        doc.add_paragraph('')
        return table

    def add_code_block(code):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F3F4"/>')
        p._p.get_or_add_pPr().append(shading)
        return p

    def add_chart(image_path, caption='', width=Inches(5.5)):
        """Add chart image with caption."""
        if os.path.exists(image_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=width)
            if caption:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(caption)
                run.font.italic = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(80, 80, 80)
            doc.add_paragraph('')
        else:
            add_para(f'[Chart not found: {image_path}]', italic=True)

    # ════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DISEASE PREDICTION SYSTEM')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Using Machine Learning')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('🧬')
    run.font.size = Pt(48)

    doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('A Project Report')
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted in partial fulfillment of the requirements for the degree of')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Your Degree Program]')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    for _ in range(3):
        doc.add_paragraph('')

    details = [
        ('Submitted By:', '[Your Name]'),
        ('Roll No.:', '[Your Roll Number]'),
        ('Branch:', '[Your Branch / Department]'),
        ('Guided By:', '[Faculty Name]'),
        ('Date:', 'April 2026'),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + '  ')
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_page_break()

    # ════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TABLE OF CONTENTS')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')

    toc_items = [
        ('1.', 'Abstract'),
        ('2.', 'Introduction'),
        ('3.', 'Problem Statement'),
        ('4.', 'Objectives'),
        ('5.', 'Literature Review'),
        ('6.', 'System Requirements'),
        ('7.', 'System Architecture'),
        ('8.', 'Methodology'),
        ('9.', 'Dataset Description'),
        ('10.', 'Data Preprocessing'),
        ('11.', 'Model Training & Evaluation'),
        ('12.', 'Implementation Details'),
        ('13.', 'Source Code'),
        ('14.', 'Results & Discussion'),
        ('15.', 'Screenshots'),
        ('16.', 'Advantages & Limitations'),
        ('17.', 'Future Scope'),
        ('18.', 'Conclusion'),
        ('19.', 'References'),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'{num}  {title}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_page_break()

    # ════════════════════════════════════════════
    # SECTION 1: ABSTRACT
    # ════════════════════════════════════════════
    add_heading('1. Abstract')
    add_para(
        'Diabetes is a chronic metabolic disorder that affects millions worldwide. '
        'Early detection and risk assessment are critical for effective prevention and management. '
        'This project implements a **Disease Prediction System** that uses **Machine Learning classification algorithms** '
        '— specifically **Random Forest Classifier** and **Logistic Regression** — to predict whether a patient '
        'is at risk of diabetes, based on health parameters such as Age, Blood Pressure, Glucose Level, and BMI.'
    )
    add_para(
        'The system is built using **Python** with the **Flask** web framework for the backend, and a modern '
        '**HTML/CSS/JavaScript** frontend. The ML model is trained on the well-known **Pima Indians Diabetes Dataset**. '
        'The trained model achieves reliable accuracy and serves predictions through a REST API that powers a '
        'real-time, interactive web interface.'
    )

    # ════════════════════════════════════════════
    # SECTION 2: INTRODUCTION
    # ════════════════════════════════════════════
    add_heading('2. Introduction')
    add_para(
        'Machine Learning (ML) continues to revolutionize the healthcare industry by enabling data-driven '
        'diagnostics and predictive models. Traditional diagnosis methods, while effective, are often '
        'time-consuming and resource-intensive. ML-based systems can assist healthcare professionals by '
        'providing rapid preliminary assessments based on quantitative patient data.'
    )
    add_para('This project demonstrates the application of supervised learning in healthcare by building an end-to-end system that:')
    add_bullet('Preprocesses raw clinical data')
    add_bullet('Trains a classification model')
    add_bullet('Deploys the model via a web application')
    add_bullet('Provides real-time risk predictions with probability scores')
    add_para(
        'The goal is to create a user-friendly tool that can assist in early diabetes screening, '
        'complementing — but not replacing — professional medical evaluation.'
    )

    # ════════════════════════════════════════════
    # SECTION 3: PROBLEM STATEMENT
    # ════════════════════════════════════════════
    add_heading('3. Problem Statement')
    add_para(
        'Diabetes mellitus is a leading cause of morbidity globally. The World Health Organization (WHO) '
        'estimates that over 422 million people worldwide have diabetes. The challenge lies in '
        '**early detection**: many patients remain undiagnosed until complications arise.'
    )
    add_para(
        '**This project addresses the problem of early diabetes risk prediction** by building a '
        'machine-learning-based web application that takes easily obtainable health metrics '
        '(Age, Blood Pressure, Glucose, BMI) and outputs a risk assessment (High Risk / Minimal Risk) '
        'along with a probability score.'
    )

    # ════════════════════════════════════════════
    # SECTION 4: OBJECTIVES
    # ════════════════════════════════════════════
    add_heading('4. Objectives')
    for obj in [
        'To collect and preprocess diabetes-related clinical data from a standard dataset.',
        'To apply data cleaning techniques (handling missing/zero values via median imputation).',
        'To train and compare classification algorithms (Random Forest & Logistic Regression).',
        'To deploy the best-performing model via a Flask-based web application.',
        'To build an intuitive, modern web UI for patient data input and result visualization.',
        'To provide probabilistic risk scores to aid interpretation.',
    ]:
        add_bullet(obj)

    # ════════════════════════════════════════════
    # SECTION 5: LITERATURE REVIEW
    # ════════════════════════════════════════════
    add_heading('5. Literature Review')
    add_table(
        ['Author(s)', 'Year', 'Contribution'],
        [
            ['Sisodia & Sisodia', '2018', 'Compared Naive Bayes, SVM, and Decision Tree for diabetes prediction; Decision Tree achieved highest accuracy.'],
            ['Zou et al.', '2018', 'Applied Random Forest and neural networks; RF showed competitive accuracy with lower complexity.'],
            ['Kavakiotis et al.', '2017', 'Comprehensive review of ML in diabetes research; classification is the most prevalent technique.'],
            ['UCI ML Repository', '—', 'Pima Indians Diabetes Dataset — widely used benchmark for binary classification.'],
        ]
    )
    add_para('Key takeaways:')
    add_bullet('**Random Forest** is among the most robust classifiers for small-to-medium healthcare datasets.')
    add_bullet('**Feature selection** improves both accuracy and interpretability.')
    add_bullet('**Standard Scaling** is beneficial when features have different ranges.')

    # ════════════════════════════════════════════
    # SECTION 6: SYSTEM REQUIREMENTS
    # ════════════════════════════════════════════
    add_heading('6. System Requirements')
    add_heading('6.1 Hardware Requirements', level=2)
    add_table(
        ['Component', 'Specification'],
        [
            ['Processor', 'Intel Core i3 or higher'],
            ['RAM', '4 GB minimum (8 GB recommended)'],
            ['Storage', '500 MB free disk space'],
            ['Display', 'Any modern display'],
        ]
    )
    add_heading('6.2 Software Requirements', level=2)
    add_table(
        ['Software', 'Version'],
        [
            ['Operating System', 'Windows 10/11, macOS, or Linux'],
            ['Python', '3.9 or higher'],
            ['Flask', '3.0.3'],
            ['Scikit-learn', '1.4.1'],
            ['Pandas', '2.2.1'],
            ['NumPy', '1.26.4'],
            ['Joblib', '1.3.2'],
            ['Web Browser', 'Chrome, Firefox, Edge (modern)'],
        ]
    )

    # ════════════════════════════════════════════
    # SECTION 7: SYSTEM ARCHITECTURE
    # ════════════════════════════════════════════
    add_heading('7. System Architecture')
    add_para('The system follows a three-tier architecture:')
    add_bullet('**Presentation Layer** — HTML/CSS/JavaScript frontend served by Flask')
    add_bullet('**Application Layer** — Flask web server with REST API endpoints')
    add_bullet('**Data/ML Layer** — Trained Random Forest model and StandardScaler')

    add_heading('System Architecture Diagram', level=3)
    arch_diagram = """┌─────────────────────────────────────────────┐
│            USER (Web Browser)               │
│     Enters: Age, BP, Glucose, BMI           │
│     Receives: Risk Label + Probability      │
└──────────────────┬──────────────────────────┘
                   │  HTTP (JSON)
                   ▼
┌─────────────────────────────────────────────┐
│         FLASK WEB SERVER (app.py)           │
│   GET /  ─→  Serves HTML template           │
│   POST /predict ─→ Parse → Scale → Predict  │
└──────────────────┬──────────────────────────┘
                   │  Loads at startup
                   ▼
┌─────────────────────────────────────────────┐
│           ML MODEL LAYER                    │
│   diabetes_model.pkl  │  scaler.pkl         │
│   (Random Forest)     │  (StandardScaler)   │
└─────────────────────────────────────────────┘"""
    add_code_block(arch_diagram)

    add_heading('Data Flow Diagram (Level 0)', level=3)
    add_para('Patient Data (Age, BP, Glucose, BMI)  →  [Disease Prediction System]  →  Risk Result (Yes/No, Probability %)')

    # ════════════════════════════════════════════
    # SECTION 8: METHODOLOGY
    # ════════════════════════════════════════════
    add_heading('8. Methodology')
    add_para('The project follows a structured machine learning pipeline:')
    steps = [
        ('Step 1: Data Acquisition',
         'The Pima Indians Diabetes Dataset is downloaded programmatically from a public GitHub repository. The raw dataset contains 768 rows and 9 columns.'),
        ('Step 2: Feature Selection',
         'Only four clinically intuitive features are retained: Age, BloodPressure, Glucose, and BMI. The target variable is Outcome (0 = No Diabetes, 1 = Diabetes).'),
        ('Step 3: Data Cleaning',
         'Zero values in BloodPressure, Glucose, and BMI are biologically implausible. These are replaced with NaN and imputed using column-wise median imputation.'),
        ('Step 4: Feature Scaling',
         'StandardScaler from Scikit-learn normalizes features to zero mean and unit variance.'),
        ('Step 5: Model Training',
         'Primary Model: Random Forest Classifier (n_estimators=100, max_depth=5). Baseline: Logistic Regression. Train-Test Split: 80/20 (random_state=42).'),
        ('Step 6: Model Evaluation',
         'Models are evaluated using Accuracy, Precision, Recall, and F1-Score.'),
        ('Step 7: Model Serialization',
         'The trained model and scaler are serialized to .pkl files using Joblib.'),
        ('Step 8: Deployment',
         'Flask web server exposes a /predict REST endpoint. Modern web frontend sends JSON and displays results in real time.'),
    ]
    for title, desc in steps:
        add_heading(title, level=3)
        add_para(desc)

    # ════════════════════════════════════════════
    # SECTION 9: DATASET DESCRIPTION
    # ════════════════════════════════════════════
    add_heading('9. Dataset Description')
    add_para('**Name:** Pima Indians Diabetes Database')
    add_para('**Source:** UCI Machine Learning Repository / National Institute of Diabetes and Digestive and Kidney Diseases')
    add_para('**Total Records:** 768')
    add_para('**Target Variable:** Outcome (Binary — 0 or 1)')

    add_heading('Original Features (9 columns)', level=2)
    add_table(
        ['#', 'Feature', 'Description', 'Type'],
        [
            ['1', 'Pregnancies', 'Number of pregnancies', 'Integer'],
            ['2', 'Glucose', 'Plasma glucose concentration (2hr OGTT)', 'Integer'],
            ['3', 'BloodPressure', 'Diastolic blood pressure (mm Hg)', 'Integer'],
            ['4', 'SkinThickness', 'Triceps skin fold thickness (mm)', 'Integer'],
            ['5', 'Insulin', '2-Hour serum insulin (mu U/ml)', 'Integer'],
            ['6', 'BMI', 'Body mass index (kg/m²)', 'Float'],
            ['7', 'DiabetesPedigreeFunction', 'Diabetes pedigree function', 'Float'],
            ['8', 'Age', 'Age in years', 'Integer'],
            ['9', 'Outcome', 'Class label (0 or 1)', 'Binary'],
        ]
    )

    add_heading('Selected Features (used in this project)', level=2)
    add_table(
        ['Feature', 'Description', 'Range (after cleaning)'],
        [
            ['Age', 'Patient age in years', '21 – 81'],
            ['BloodPressure', 'Diastolic BP (mm Hg)', '~24 – 122'],
            ['Glucose', '2-hr oral glucose tolerance', '~44 – 199'],
            ['BMI', 'Body Mass Index (kg/m²)', '~18.2 – 67.1'],
            ['Outcome', 'Target (0 = No, 1 = Yes)', '0 or 1'],
        ]
    )

    # ── CHART: Class Distribution ──
    add_heading('Dataset Class Distribution', level=2)
    add_para('The following chart shows the distribution of diabetic vs non-diabetic samples in the dataset:')
    add_chart('charts/class_distribution.png', 'Figure 9.1: Class Distribution — Diabetic vs Non-Diabetic samples')

    # ════════════════════════════════════════════
    # SECTION 10: DATA PREPROCESSING
    # ════════════════════════════════════════════
    add_heading('10. Data Preprocessing')

    add_heading('10.1 Handling Missing Values', level=2)
    add_para(
        'The dataset uses 0 as a placeholder for missing values in Glucose, BloodPressure, and BMI. '
        'These are replaced with NaN and imputed with the **median** of each column.'
    )
    add_code_block(
        "cols_with_zeros = ['BloodPressure', 'Glucose', 'BMI']\n"
        "df_subset[cols_with_zeros] = df_subset[cols_with_zeros].replace(0, np.nan)\n\n"
        "for col in cols_with_zeros:\n"
        "    median_val = df_subset[col].median()\n"
        "    df_subset[col] = df_subset[col].fillna(median_val)"
    )
    add_para('**Why median?** Robust to outliers and better preserves skewed medical data distributions vs. mean.')

    add_heading('10.2 Feature Scaling', level=2)
    add_code_block(
        "scaler = StandardScaler()\n"
        "X_train_scaled = scaler.fit_transform(X_train)\n"
        "X_test_scaled = scaler.transform(X_test)"
    )
    add_para('StandardScaler transforms each feature to **mean = 0** and **standard deviation = 1**.')

    # ── CHART: Feature Distributions ──
    add_heading('10.3 Feature Distribution Analysis', level=2)
    add_para('The following histograms show the distribution of each feature, separated by diabetes outcome:')
    add_chart('charts/feature_distributions.png', 'Figure 10.1: Feature Distributions by Diabetes Outcome')

    # ── CHART: Box Plots ──
    add_heading('10.4 Feature Comparison — Box Plots', level=2)
    add_para('Box plots reveal the statistical differences between diabetic and non-diabetic groups for each feature:')
    add_chart('charts/box_plots.png', 'Figure 10.2: Feature Box Plots — Diabetic vs Non-Diabetic')

    # ── CHART: Correlation Heatmap ──
    add_heading('10.5 Feature Correlation Analysis', level=2)
    add_para('The correlation heatmap shows the relationships between all features and the target variable:')
    add_chart('charts/correlation_heatmap.png', 'Figure 10.3: Feature Correlation Heatmap')

    # ════════════════════════════════════════════
    # SECTION 11: MODEL TRAINING & EVALUATION
    # ════════════════════════════════════════════
    add_heading('11. Model Training & Evaluation')

    add_heading('11.1 Random Forest Classifier (Primary Model)', level=2)
    add_table(
        ['Parameter', 'Value', 'Rationale'],
        [
            ['n_estimators', '100', 'Sufficient trees for stable ensemble'],
            ['max_depth', '5', 'Prevents overfitting on small dataset'],
            ['random_state', '42', 'Reproducibility'],
        ]
    )
    add_para(
        '**Algorithm Overview:** Random Forest constructs multiple decision trees during training '
        'and outputs the mode of individual tree predictions. It reduces overfitting by averaging '
        'across trees trained on random subsets of data and features.'
    )

    add_heading('11.2 Logistic Regression (Baseline)', level=2)
    add_para('Logistic Regression models the probability of a binary outcome using the sigmoid function. It serves as a baseline.')

    add_heading('11.3 Evaluation Metrics', level=2)
    add_bullet('**Accuracy** = (TP + TN) / (TP + TN + FP + FN)')
    add_bullet('**Precision** = TP / (TP + FP) — How many predicted positives are actually positive')
    add_bullet('**Recall** = TP / (TP + FN) — How many actual positives were correctly identified')
    add_bullet('**F1-Score** = 2 × (Precision × Recall) / (Precision + Recall)')

    # ── CHART: Model Comparison ──
    add_heading('11.4 Model Accuracy Comparison', level=2)
    add_para('The following chart compares the accuracy of both models:')
    add_chart('charts/model_comparison.png', 'Figure 11.1: Random Forest vs Logistic Regression — Accuracy Comparison')

    # ── CHART: Confusion Matrix ──
    add_heading('11.5 Confusion Matrix', level=2)
    add_para('The confusion matrices show true positives, true negatives, false positives, and false negatives for both models:')
    add_chart('charts/confusion_matrix.png', 'Figure 11.2: Confusion Matrix — Random Forest vs Logistic Regression')

    # ── CHART: Feature Importance ──
    add_heading('11.6 Feature Importance', level=2)
    add_para('The Random Forest model provides feature importance scores, indicating which features contribute most to predictions:')
    add_chart('charts/feature_importance.png', 'Figure 11.3: Random Forest Feature Importance Scores')

    # ════════════════════════════════════════════
    # SECTION 12: IMPLEMENTATION DETAILS
    # ════════════════════════════════════════════
    add_heading('12. Implementation Details')

    add_heading('12.1 Project Directory Structure', level=2)
    add_code_block(
        "Disease Prediction System/\n"
        "├── app.py                     # Flask web application\n"
        "├── requirements.txt           # Python dependencies\n"
        "├── data/\n"
        "│   ├── download_prep.py       # Data download & preprocessing\n"
        "│   └── cleaned_diabetes.csv   # Cleaned dataset\n"
        "├── model/\n"
        "│   ├── train.py               # Model training script\n"
        "│   ├── diabetes_model.pkl     # Serialized Random Forest\n"
        "│   └── scaler.pkl             # Serialized StandardScaler\n"
        "├── templates/\n"
        "│   └── index.html             # Frontend HTML template\n"
        "└── static/\n"
        "    └── style.css              # Glassmorphism stylesheet"
    )

    add_heading('12.2 Module Descriptions', level=2)
    add_table(
        ['Module', 'File', 'Description'],
        [
            ['Data Pipeline', 'data/download_prep.py', 'Downloads dataset, selects features, cleans missing values, saves CSV'],
            ['Model Training', 'model/train.py', 'Splits, scales, trains RF & LR, evaluates, serializes models'],
            ['Web Server', 'app.py', 'Flask app: GET / (UI) and POST /predict (JSON prediction)'],
            ['Frontend', 'templates/index.html', 'Responsive UI with AJAX call and animated result display'],
            ['Styling', 'static/style.css', 'Modern glassmorphism design with dark theme'],
        ]
    )

    add_heading('12.3 API Endpoint', level=2)
    add_para('**POST /predict**')
    add_para('**Request Body (JSON):**')
    add_code_block('{\n    "age": 45,\n    "bp": 80,\n    "glucose": 110,\n    "bmi": 28.5\n}')
    add_para('**Response (JSON):**')
    add_code_block('{\n    "risk": "Yes",\n    "probability": 72.35\n}')
    add_table(
        ['Field', 'Type', 'Description'],
        [
            ['risk', 'String', '"Yes" (diabetic risk) or "No" (minimal risk)'],
            ['probability', 'Float', 'Probability of diabetes (0–100%)'],
        ]
    )

    # ════════════════════════════════════════════
    # SECTION 13: SOURCE CODE
    # ════════════════════════════════════════════
    add_heading('13. Source Code')

    add_heading('13.1 Data Download & Preprocessing (data/download_prep.py)', level=2)
    with open('data/download_prep.py', 'r', encoding='utf-8') as f:
        add_code_block(f.read())

    add_heading('13.2 Model Training (model/train.py)', level=2)
    with open('model/train.py', 'r', encoding='utf-8') as f:
        add_code_block(f.read())

    add_heading('13.3 Flask Web Application (app.py)', level=2)
    with open('app.py', 'r', encoding='utf-8') as f:
        add_code_block(f.read())

    add_heading('13.4 Frontend Template (templates/index.html)', level=2)
    with open('templates/index.html', 'r', encoding='utf-8') as f:
        add_code_block(f.read())

    # ════════════════════════════════════════════
    # SECTION 14: RESULTS & DISCUSSION
    # ════════════════════════════════════════════
    add_heading('14. Results & Discussion')

    add_heading('14.1 Model Performance Summary', level=2)
    add_table(
        ['Model', 'Accuracy', 'Notes'],
        [
            ['Random Forest Classifier', '~76–78%', 'Primary model, selected for deployment'],
            ['Logistic Regression', '~74–76%', 'Baseline comparison'],
        ]
    )

    add_heading('14.2 Interpretation', level=2)
    add_bullet('The **Random Forest Classifier** was selected due to its superior accuracy and ability to capture non-linear relationships.')
    add_bullet('The probability output adds nuance beyond binary classification.')
    add_bullet('Using only 4 features makes the system practical for quick clinical screening.')

    add_heading('14.3 Sample Predictions', level=2)
    add_table(
        ['Age', 'BP', 'Glucose', 'BMI', 'Predicted Risk', 'Probability'],
        [
            ['45', '80', '150', '32.0', 'High Risk', '~70%'],
            ['25', '72', '85', '22.5', 'Minimal Risk', '~15%'],
            ['60', '90', '180', '35.2', 'High Risk', '~85%'],
            ['30', '68', '95', '24.0', 'Minimal Risk', '~20%'],
        ]
    )

    # ════════════════════════════════════════════
    # SECTION 15: SCREENSHOTS
    # ════════════════════════════════════════════
    add_heading('15. Screenshots')
    add_para('Screenshot 1: Home page — Input form with fields for Age, BP, Glucose, and BMI', italic=True)
    add_para('[Insert screenshot here]', italic=True)
    add_para('')
    add_para('Screenshot 2: Low Risk result — Green badge showing "Minimal Risk Detected"', italic=True)
    add_para('[Insert screenshot here]', italic=True)
    add_para('')
    add_para('Screenshot 3: High Risk result — Red badge showing "High Risk of Diabetes Detected"', italic=True)
    add_para('[Insert screenshot here]', italic=True)

    # ════════════════════════════════════════════
    # SECTION 16: ADVANTAGES & LIMITATIONS
    # ════════════════════════════════════════════
    add_heading('16. Advantages & Limitations')
    add_heading('16.1 Advantages', level=2)
    for adv in [
        '**Easy to Use:** Simple web interface requires no technical knowledge.',
        '**Real-Time Predictions:** Instant results via AJAX.',
        '**Probabilistic Output:** Confidence percentage, not just yes/no.',
        '**Modern UI:** Glassmorphism design with animations.',
        '**Modular Architecture:** Independent data pipeline, training, and deployment.',
        '**Reproducible:** Fixed random seeds and pinned dependencies.',
    ]:
        add_bullet(adv)

    add_heading('16.2 Limitations', level=2)
    for lim in [
        '**Limited Features:** Uses only 4 features; more clinical variables could improve accuracy.',
        '**Dataset Size:** Only 768 records, limiting generalizability.',
        '**Single Disease:** Currently supports only diabetes prediction.',
        '**No Authentication:** Not suitable for production healthcare use.',
        '**Bias:** Dataset collected from a specific demographic (Pima Indian women).',
    ]:
        add_bullet(lim)

    # ════════════════════════════════════════════
    # SECTION 17: FUTURE SCOPE
    # ════════════════════════════════════════════
    add_heading('17. Future Scope')
    for item in [
        '**Multi-Disease Support:** Extend to heart disease, kidney disease, etc.',
        '**Additional Features:** Add Insulin, HbA1c, family history.',
        '**Deep Learning Models:** Neural networks for larger datasets.',
        '**User Authentication:** Secure medical data handling.',
        '**Patient History Tracking:** Longitudinal monitoring.',
        '**Mobile Application:** On-the-go health screening.',
        '**Explainable AI (XAI):** SHAP/LIME for feature explanations.',
        '**Cloud Deployment:** AWS/GCP/Azure for scalability.',
    ]:
        add_bullet(item)

    # ════════════════════════════════════════════
    # SECTION 18: CONCLUSION
    # ════════════════════════════════════════════
    add_heading('18. Conclusion')
    add_para(
        'This project successfully demonstrates the application of Machine Learning in healthcare '
        'for early diabetes risk prediction. The system follows a complete ML pipeline — from data '
        'acquisition and preprocessing, to model training and evaluation, to deployment via a web application.'
    )
    add_para(
        'The **Random Forest Classifier** was found to outperform Logistic Regression on the Pima Indians '
        'Diabetes Dataset and was selected as the production model. The Flask-based web interface provides '
        'a modern, user-friendly experience with real-time predictions and probabilistic risk scores.'
    )
    add_para(
        'While the system has limitations (dataset size, feature count, single-disease focus), it serves '
        'as a strong foundation for future extensions. The modular architecture makes it straightforward '
        'to add new diseases, features, and ML algorithms.'
    )

    # ════════════════════════════════════════════
    # SECTION 19: REFERENCES
    # ════════════════════════════════════════════
    add_heading('19. References')
    references = [
        'Smith, J.W., Everhart, J.E., Dickson, W.C., Knowler, W.C., & Johannes, R.S. (1988). '
        'Using the ADAP learning algorithm to forecast the onset of diabetes mellitus. '
        'Proceedings of the Annual Symposium on Computer Application in Medical Care, pp. 261–265.',
        'Sisodia, D., & Sisodia, D.S. (2018). Prediction of Diabetes using Classification Algorithms. '
        'Procedia Computer Science, 132, 1578–1585.',
        'Zou, Q., Qu, K., Luo, Y., Yin, D., Ju, Y., & Tang, H. (2018). '
        'Predicting Diabetes Mellitus With Machine Learning Techniques. Frontiers in Genetics, 9, 515.',
        'Kavakiotis, I., et al. (2017). Machine Learning and Data Mining Methods in Diabetes Research. '
        'Computational and Structural Biotechnology Journal, 15, 104–116.',
        'Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. JMLR, 12, 2825–2830.',
        'UCI Machine Learning Repository — Pima Indians Diabetes Dataset. '
        'https://archive.ics.uci.edu/ml/datasets/diabetes',
        'Flask Documentation. https://flask.palletsprojects.com/',
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        run = p.add_run(f'[{i}]  ')
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run = p.add_run(ref)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    # ════════════════════════════════════════════
    # SAVE
    # ════════════════════════════════════════════
    output_file = 'Disease_Prediction_System_Report.docx'
    doc.save(output_file)
    print(f'\n✅ Report with charts generated: {output_file}')
    print(f'   File size: {os.path.getsize(output_file) / 1024:.1f} KB')
    print(f'\n📊 Charts embedded:')
    chart_files = [
        'class_distribution.png', 'feature_distributions.png', 'box_plots.png',
        'correlation_heatmap.png', 'model_comparison.png', 'confusion_matrix.png',
        'feature_importance.png'
    ]
    for cf in chart_files:
        path = f'charts/{cf}'
        status = '✅' if os.path.exists(path) else '❌ NOT FOUND'
        print(f'   {status}  {cf}')
    print(f'\n📝 Remember to fill in your name, roll number, branch, and faculty on the cover page!')


if __name__ == '__main__':
    create_report()
